"""DOCX parser."""

from __future__ import annotations

from pathlib import Path

import docx

from ..models import ParsedDocument, ParsedSection


def _heading_level(style_name: str) -> int | None:
    if not style_name.lower().startswith("heading"):
        return None
    digits = "".join(ch for ch in style_name if ch.isdigit())
    if not digits:
        return 1
    try:
        return max(1, int(digits))
    except ValueError:
        return 1


def parse_docx(path: Path) -> ParsedDocument:
    document = docx.Document(str(path))
    sections: list[ParsedSection] = []
    breadcrumb_stack: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        level = _heading_level(style_name or "")
        if level is not None:
            while len(breadcrumb_stack) >= level:
                breadcrumb_stack.pop()
            breadcrumb_stack.append(text)
            continue
        heading = breadcrumb_stack[-1] if breadcrumb_stack else None
        sections.append(
            ParsedSection(text=text, heading=heading, breadcrumbs=breadcrumb_stack.copy())
        )

    return ParsedDocument(
        source_path=path,
        source_type="docx",
        sections=sections,
    )
